"""
Document Reader for RAG System
Supports PDF, DOCX, and TXT files
"""

import logging
from pathlib import Path
from typing import Tuple, Dict, Any, Optional
import mimetypes
from datetime import datetime

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf
        import PyPDF2
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentReader:
    """
    Universal document reader for business documents.
    Supports PDF, DOCX, and TXT files with metadata extraction.
    """
    
    def __init__(self):
        """Initialize the document reader."""
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        logger.info("DocumentReader initialized")
        
        # Log available readers
        if not PDF_AVAILABLE:
            logger.warning("PDF support not available. Install PyPDF2: pip install PyPDF2")
        if not DOCX_AVAILABLE:
            logger.warning("DOCX support not available. Install python-docx: pip install python-docx")
    
    def read_document(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Read a document and return its content with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (content_string, metadata_dict)
            
        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
            Exception: If reading fails
        """
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Check if file type is supported
        extension = file_path.suffix.lower()
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {self.supported_extensions}")
        
        logger.info(f"Reading document: {file_path}")
        
        try:
            # Extract basic metadata
            metadata = self._extract_basic_metadata(file_path)
            
            # Read content based on file type
            if extension == '.pdf':
                content = self._read_pdf(file_path, metadata)
            elif extension == '.docx':
                content = self._read_docx(file_path, metadata)
            elif extension == '.txt':
                content = self._read_txt(file_path, metadata)
            else:
                raise ValueError(f"Unsupported file extension: {extension}")
            
            # Validate content
            if not content or not content.strip():
                logger.warning(f"No content extracted from {file_path}")
                content = f"[No content could be extracted from {file_path.name}]"
            
            # Add content statistics to metadata
            metadata.update(self._analyze_content(content))
            
            logger.info(f"Successfully read {file_path}: {len(content)} characters, {metadata['word_count']} words")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            raise
    
    def _extract_basic_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic file metadata."""
        stat = file_path.stat()
        
        return {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_extension': file_path.suffix.lower(),
            'file_size_bytes': stat.st_size,
            'created_date': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_date': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'content_type': mimetypes.guess_type(str(file_path))[0] or 'unknown'
        }
    
    def _read_pdf(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Read PDF file content."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF support not available. Install PyPDF2: pip install PyPDF2")
        
        try:
            content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Add PDF-specific metadata
                if pdf_reader.metadata:
                    pdf_metadata = pdf_reader.metadata
                    metadata.update({
                        'pdf_title': pdf_metadata.get('/Title', '').strip(),
                        'pdf_author': pdf_metadata.get('/Author', '').strip(),
                        'pdf_subject': pdf_metadata.get('/Subject', '').strip(),
                        'pdf_creator': pdf_metadata.get('/Creator', '').strip(),
                        'pdf_producer': pdf_metadata.get('/Producer', '').strip(),
                    })
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            content.append(f"\n--- Page {page_num} ---\n")
                            content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {e}")
                        content.append(f"\n--- Page {page_num} (extraction failed) ---\n")
            
            return '\n'.join(content).strip()
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise
    
    def _read_docx(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Read DOCX file content."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX support not available. Install python-docx: pip install python-docx")
        
        try:
            doc = Document(file_path)
            content = []
            
            # Extract document properties
            if doc.core_properties:
                props = doc.core_properties
                metadata.update({
                    'docx_title': props.title or '',
                    'docx_author': props.author or '',
                    'docx_subject': props.subject or '',
                    'docx_category': props.category or '',
                    'docx_comments': props.comments or '',
                    'docx_created': props.created.isoformat() if props.created else '',
                    'docx_modified': props.modified.isoformat() if props.modified else '',
                })
            
            # Extract paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
                    paragraph_count += 1
            
            metadata['paragraph_count'] = paragraph_count
            
            # Extract tables if any
            table_count = len(doc.tables)
            if table_count > 0:
                metadata['table_count'] = table_count
                
                for table_num, table in enumerate(doc.tables, 1):
                    content.append(f"\n--- Table {table_num} ---\n")
                    
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_cells.append(cell_text)
                        if row_cells:
                            content.append(" | ".join(row_cells))
            
            return '\n'.join(content).strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise
    
    def _read_txt(self, file_path: Path, metadata: Dict[str, Any]) -> str:
        """Read TXT file content."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    metadata['text_encoding'] = encoding
                    return content.strip()
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try binary read
            logger.warning(f"Could not decode {file_path} with standard encodings, trying binary read")
            
            with open(file_path, 'rb') as file:
                raw_content = file.read()
            
            # Try to decode as UTF-8 with error handling
            content = raw_content.decode('utf-8', errors='replace')
            metadata['text_encoding'] = 'utf-8-with-errors'
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            raise
    
    def _analyze_content(self, content: str) -> Dict[str, Any]:
        """Analyze content and return statistics."""
        lines = content.split('\n')
        words = content.split()
        
        return {
            'character_count': len(content),
            'word_count': len(words),
            'line_count': len(lines),
            'non_empty_line_count': len([line for line in lines if line.strip()]),
            'average_words_per_line': len(words) / max(len([line for line in lines if line.strip()]), 1)
        }
    
    def get_supported_extensions(self) -> set:
        """Get set of supported file extensions."""
        return self.supported_extensions.copy()
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if a file is supported."""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def scan_directory(self, directory: Path, recursive: bool = False) -> list[Path]:
        """
        Scan a directory for supported documents.
        
        Args:
            directory: Directory to scan
            recursive: Whether to scan subdirectories
            
        Returns:
            List of supported document paths
        """
        directory = Path(directory)
        
        if not directory.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return []
        
        if not directory.is_dir():
            logger.warning(f"Path is not a directory: {directory}")
            return []
        
        documents = []
        
        # Choose scanning method
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        # Find all files with supported extensions
        for ext in self.supported_extensions:
            pattern_with_ext = pattern + ext
            found_files = directory.glob(pattern_with_ext)
            documents.extend(found_files)
        
        # Sort by name for consistent ordering
        documents.sort(key=lambda x: x.name.lower())
        
        logger.info(f"Found {len(documents)} supported documents in {directory}")
        return documents