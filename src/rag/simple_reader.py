"""
Step 1: Simple Document Reader
Just extract text from PDF, DOCX, and TXT files
"""
import os
from pathlib import Path
from typing import Optional

import PyPDF2
from docx import Document


class SimpleDocumentReader:
    """Basic document reader - just extract text, nothing fancy yet"""
    
    def __init__(self, max_file_size_mb: int = 50):
        """
        Initialize reader with file size limit
        
        Args:
            max_file_size_mb: Maximum file size to process in MB
        """
        self.max_file_size_mb = max_file_size_mb
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
    
    def _check_file_size(self, file_path: str) -> bool:
        """Check if file is within size limits"""
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size_bytes:
            print(f"Warning: File {file_path} is too large ({file_size / (1024*1024):.1f}MB)")
            return False
        return True
    
    def read_pdf(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text or None if failed
        """
        if not self._check_file_size(file_path):
            return None
            
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                print(f"Reading PDF with {len(pdf_reader.pages)} pages...")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"  Page {page_num}: {len(page_text)} characters")
                
                return text
                
        except Exception as e:
            print(f"Error reading PDF {file_path}: {str(e)}")
            return None
    
    def read_docx(self, file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None if failed
        """
        if not self._check_file_size(file_path):
            return None
            
        try:
            doc = Document(file_path)
            text = ""
            
            print(f"Reading DOCX with {len(doc.paragraphs)} paragraphs...")
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text += paragraph.text + "\n"
                if para_num <= 5:  # Show first 5 paragraphs info
                    print(f"  Paragraph {para_num}: {len(paragraph.text)} characters")
            
            return text
            
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {str(e)}")
            return None
    
    def read_txt(self, file_path: str) -> Optional[str]:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text or None if failed
        """
        if not self._check_file_size(file_path):
            return None
            
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                print(f"Reading TXT file: {len(text)} characters")
                return text
                
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    print(f"Reading TXT file (latin-1 encoding): {len(text)} characters")
                    return text
            except Exception as e:
                print(f"Error reading TXT {file_path}: {str(e)}")
                return None
        except Exception as e:
            print(f"Error reading TXT {file_path}: {str(e)}")
            return None
    
    def read_document(self, file_path: str) -> Optional[str]:
        """
        Read any supported document type
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text or None if failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"File not found: {file_path}")
            return None
        
        extension = file_path.suffix.lower()
        
        print(f"\n--- Reading {file_path.name} ---")
        
        if extension == '.pdf':
            return self.read_pdf(str(file_path))
        elif extension == '.docx':
            return self.read_docx(str(file_path))
        elif extension == '.txt':
            return self.read_txt(str(file_path))
        else:
            print(f"Unsupported file type: {extension}")
            return None


# Test the reader
if __name__ == "__main__":
    # Initialize reader
    reader = SimpleDocumentReader()
    
    # Test with a sample file (you'll need to provide your own)
    test_file = "documents/sample_contract.pdf"  # Replace with your file
    
    if Path(test_file).exists():
        text = reader.read_document(test_file)
        if text:
            print(f"\n--- Extracted Text Preview ---")
            print(f"Total length: {len(text)} characters")
            print(f"First 500 characters:")
            print("-" * 50)
            print(text[:500])
            print("-" * 50)
        else:
            print("Failed to extract text")
    else:
        print(f"Please add a test document to: {test_file}")
        print("Supported formats: PDF, DOCX, TXT")
        
        # Show what files are in documents folder
        docs_folder = Path("documents")
        if docs_folder.exists():
            files = list(docs_folder.glob("*"))
            if files:
                print(f"\nFiles found in documents folder:")
                for file in files:
                    print(f"  - {file.name}")
            else:
                print(f"\nNo files found in documents folder")
        else:
            print(f"\nPlease create documents folder and add some test files")